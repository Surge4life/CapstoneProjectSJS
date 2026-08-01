"""
Policy-to-Code engine — the UDOC governance protocol layer.

Pipeline (assistive + human-in-the-loop, by design):
  1) extract_text(): pull text from an uploaded PDF / DOCX / TXT of passed legislation.
  2) extract_rules(): compile candidate machine-enforceable rules from the text using
     transparent, auditable heuristics. Every rule keeps the exact source excerpt and a
     confidence score, and is editable. NOTHING is enforced until a human ACTIVATES the pack.
  3) apply(): at decision time, the active rule set is evaluated against the decision context
     and can BLOCK / downgrade to REVIEW / FLAG — making uploaded legislation operationally
     enforced inside the non-bypassable EVA path.

This is deliberately explainable rather than a black-box legal NLP: a compliance officer can
see exactly which clause produced which rule, edit it, and approve it before it governs.
"""
import io
import re
import time
from typing import List, Dict, Any
from sqlalchemy import select
from sqlalchemy.orm import Session
from app.db.models import PolicyPack, PolicyRule
from app.services.governance_bridge import RISK_TIER_R


# ─────────────────────────── text extraction ───────────────────────────
def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            r = PdfReader(io.BytesIO(data))
            return "\n".join((p.extract_text() or "") for p in r.pages)
        if name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for t in doc.tables:
                for row in t.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        return data.decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[extraction-error: {e}]"


def _sentences(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", text or "")
    raw = re.split(r"(?<=[.;:])\s+|\n+", text)
    return [s.strip() for s in raw if 12 <= len(s.strip()) <= 400]


_STOP = set("the a an of to and or for in on at by with from as is are be that this which shall must "
            "not no any all such where when who whom whose will may can it its their his her they we you".split())


def _keywords(s: str, n: int = 6) -> str:
    words = re.findall(r"[A-Za-z][A-Za-z\-]{3,}", s.lower())
    seen, out = set(), []
    for w in words:
        if w in _STOP or w in seen:
            continue
        seen.add(w); out.append(w)
        if len(out) >= n:
            break
    return " ".join(out)


# pattern → (kind, severity, operator, threshold, description)
_RULES = [
    (r"\b(prohibit|prohibited|shall not|must not|may not|banned?|not permitted|forbidden|unacceptable risk)\b",
     "PROHIBIT", "BLOCK", "contains", None, "Prohibited practice — systems matching this are blocked."),
    (r"\b(human (oversight|review|in[\- ]the[\- ]loop|intervention)|meaningful human|human agency)\b",
     "REQUIRE_HITL", "REVIEW", ">=", None, "Human oversight required — high-risk decisions routed to review."),
    (r"\b(high[\- ]risk)\b",
     "RISK_TIER_CAP", "REVIEW", ">=", 0.8, "High-risk systems require conformity review before approval."),
    (r"\b(discriminat|bias|fairness|equitab|protected (group|characteristic|attribute)|disparate)\b",
     "MAX_DISPARATE_IMPACT", "REVIEW", ">=", 0.8, "Non-discrimination — fairness (disparate-impact) floor enforced."),
    (r"\b(data (localis|localiz|residency|sovereignty)|stored within|cross[\- ]border|personal information|popia)\b",
     "DATA_LOCALISATION", "REVIEW", ">=", 1.0, "Data localisation / sovereignty of personal information."),
    (r"\b(sovereign|jurisdiction|within (the )?republic|south africa|national borders)\b",
     "MIN_SOVEREIGNTY", "REVIEW", ">=", 1.0, "Sovereignty — decisioning must execute under in-jurisdiction signals."),
    (r"\b(transparen|disclose|inform(ed)?|explainab|notice|right to (an )?explanation)\b",
     "KEYWORD_FLAG", "FLAG", "contains", None, "Transparency / disclosure obligation — flagged for evidence."),
    (r"\b(record[\- ]keeping|logging|traceab|audit|register(ed|ation)?)\b",
     "KEYWORD_FLAG", "FLAG", "contains", None, "Record-keeping / audit obligation — flagged for evidence."),
]


def extract_rules(text: str) -> List[Dict]:
    """Compile transparent candidate rules from legislation text. Human reviews before activation."""
    out: List[Dict] = []
    seen = set()
    n = 0
    for s in _sentences(text):
        low = s.lower()
        for pat, kind, severity, op, thr, desc in _RULES:
            if re.search(pat, low):
                key = (kind, _keywords(s))
                if key in seen:
                    continue
                seen.add(key)
                n += 1
                out.append({
                    "code": f"PR-{n:03d}", "kind": kind, "severity": severity,
                    "operator": op, "threshold": thr,
                    "target": _keywords(s, 8),
                    "description": desc,
                    "source_excerpt": s[:380],
                    "confidence": round(0.55 + 0.05 * min(4, low.count(" ") // 8), 2),
                    "enabled": True,
                })
                break
        if n >= 60:
            break
    return out


def summarise(text: str, rules: List[Dict]) -> str:
    kinds: Dict[str, int] = {}
    for r in rules:
        kinds[r["kind"]] = kinds.get(r["kind"], 0) + 1
    head = (text or "").strip().split("\n")[0][:140]
    bits = ", ".join(f"{k}×{v}" for k, v in sorted(kinds.items()))
    return f"{len(rules)} candidate rules compiled ({bits or 'none'}). Opening: “{head}”."


# ─────────────────────────── enforcement ───────────────────────────
# Cache PLAIN DICTS only — never ORM instances (DetachedInstanceError across requests).
_EPOCH = 0
_MEMO: Dict = {}
_LAST_RELOAD_MS = 0.0


def invalidate():
    """Bump the policy epoch so the next decision hot-reloads the active rule set."""
    global _EPOCH, _MEMO
    _EPOCH += 1
    _MEMO = {}


def hot_reload_stats() -> Dict:
    return {"epoch": _EPOCH, "last_reload_ms": round(_LAST_RELOAD_MS, 3), "cached_keys": len(_MEMO)}


def _rule_dict(r: PolicyRule) -> Dict[str, Any]:
    return {
        "code": r.code or "",
        "kind": r.kind or "",
        "severity": r.severity or "REVIEW",
        "operator": r.operator or "",
        "threshold": r.threshold,
        "target": r.target or "",
        "description": r.description or "",
        "source_excerpt": r.source_excerpt or "",
        "confidence": r.confidence,
        "enabled": bool(r.enabled),
    }


def active_rules(db: Session, jurisdiction: str = None, sector: str = None, tenant_pk: int = None) -> List[Dict[str, Any]]:
    global _LAST_RELOAD_MS
    key = (_EPOCH, jurisdiction, sector, tenant_pk)
    hit = _MEMO.get(key)
    if hit is not None:
        return hit
    _t0 = time.perf_counter()
    packs = db.execute(select(PolicyPack).where(PolicyPack.status == "ACTIVE")).scalars().all()
    pack_ids = []
    for p in packs:
        if tenant_pk is not None and p.tenant_pk is not None and p.tenant_pk != tenant_pk:
            continue
        if jurisdiction and p.jurisdiction not in (jurisdiction, "GLOBAL", "*"):
            continue
        if sector and p.sector not in (sector, "GENERAL"):
            continue
        pack_ids.append(p.id)
    if not pack_ids:
        rows: List[Dict[str, Any]] = []
    else:
        orm_rows = db.execute(select(PolicyRule).where(
            PolicyRule.pack_id.in_(pack_ids), PolicyRule.enabled == True)).scalars().all()  # noqa: E712
        rows = [_rule_dict(r) for r in orm_rows]
    _MEMO[key] = rows
    _LAST_RELOAD_MS = (time.perf_counter() - _t0) * 1000.0
    return rows


def apply(db: Session, ev, model, verdict) -> Dict:
    """Evaluate active policy rules against a decision; return findings + adjusted decision."""
    rules = active_rules(db, jurisdiction=getattr(model, "jurisdiction", "ZA"),
                         sector=None, tenant_pk=getattr(model, "tenant_pk", None))
    findings = []
    risk_num = RISK_TIER_R.get(ev.risk_tier, 0.5)
    haystack = f"{getattr(model,'use_case','')} {getattr(model,'name','')} {ev.model_id}".lower()

    for r in rules:
        fired, msg = False, ""
        kind = r.get("kind") or ""
        thr = r.get("threshold")
        sev = r.get("severity") or "REVIEW"
        target = r.get("target") or ""
        code = r.get("code") or ""
        excerpt = (r.get("source_excerpt") or "")[:160]

        if kind == "PROHIBIT":
            kws = [k for k in target.split() if len(k) > 3]
            hit = [k for k in kws if k in haystack]
            if hit or ev.risk_tier == "UNACCEPTABLE":
                fired = True; msg = f"prohibited practice (matched: {', '.join(hit) or ev.risk_tier})"
        elif kind == "RISK_TIER_CAP":
            if thr is not None and risk_num >= thr:
                fired = True; msg = f"risk {risk_num:.2f} ≥ cap {thr:.2f}"
        elif kind == "REQUIRE_HITL":
            if ev.risk_tier in ("HIGH", "UNACCEPTABLE"):
                fired = True; msg = "high-risk system requires human oversight"
        elif kind == "MIN_SOVEREIGNTY":
            need = thr if thr is not None else 1.0
            if verdict.sovereign_svs < need:
                fired = True; msg = f"sovereignty {verdict.sovereign_svs:.2f} < {need:.2f}"
        elif kind == "MIN_COMPLIANCE":
            need = thr if thr is not None else 0.7
            if verdict.compliance < need:
                fired = True; msg = f"compliance {verdict.compliance:.2f} < {need:.2f}"
        elif kind == "MAX_DISPARATE_IMPACT":
            need = thr if thr is not None else 0.8
            if verdict.disparate_impact < need:
                fired = True; msg = f"disparate-impact {verdict.disparate_impact:.2f} < {need:.2f}"
        elif kind == "DATA_LOCALISATION":
            if getattr(model, "jurisdiction", "ZA") != "ZA" or ev.storage < 1.0:
                fired = True; msg = "data-localisation signal not satisfied"
        elif kind == "KEYWORD_FLAG":
            fired = True; msg = "obligation flagged for evidence"

        findings.append({"code": code, "kind": kind, "severity": sev,
                         "fired": fired, "message": msg, "excerpt": excerpt})

    blockers = [f for f in findings if f["fired"] and f["severity"] == "BLOCK"]
    reviewers = [f for f in findings if f["fired"] and f["severity"] == "REVIEW"]
    decision = verdict.decision
    if blockers:
        decision = "BLOCK"
    elif reviewers and decision in ("APPROVE", "RESTRICT"):
        decision = "REVIEW"
    # EVA engine may already BLOCK on DI; keep the stronger outcome
    if decision == "BLOCK" or verdict.decision == "BLOCK":
        decision = "BLOCK"
    return {
        "adjusted_decision": decision,
        "policy_enforced": bool(rules),
        "rules_evaluated": len(rules),
        "fired": [f for f in findings if f["fired"]],
        "findings": findings,
    }
